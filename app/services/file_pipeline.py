import os
from pathlib import Path

from app.core import logger, settings
from app.services.tabular_pipeline import process_tabular
from app.services.text_pipeline import depersonalize_text


def _extract_text_from_txt(file_path: str) -> str:
    with open(file_path, "r", encoding="utf-8", errors="replace") as f:
        return f.read()


def _extract_text_from_pdf(file_path: str) -> str:
    from PyPDF2 import PdfReader

    reader = PdfReader(file_path)
    pages = []
    for page in reader.pages:
        text = page.extract_text()
        if text:
            pages.append(text)
    return "\n\n".join(pages)


def _extract_text_from_docx(file_path: str) -> str:
    from docx import Document

    doc = Document(file_path)
    return "\n\n".join(p.text for p in doc.paragraphs if p.text)


def _save_text_to_file(text: str, original_path: str) -> str:
    ext = Path(original_path).suffix.lower()
    out = original_path.replace(ext, f"_depersonalized{ext}")

    if ext == ".pdf":
        try:
            from fpdf import FPDF

            pdf = FPDF()
            pdf.add_page()
            pdf.set_font("Helvetica", size=10)
            for line in text.split("\n"):
                pdf.cell(0, 6, line, new_x="LMARGIN", new_y="NEXT")
            pdf.output(out)
        except ImportError:
            out = out.replace(".pdf", ".txt")
            with open(out, "w", encoding="utf-8") as f:
                f.write(text)
    elif ext == ".docx":
        try:
            from docx import Document

            doc = Document()
            for line in text.split("\n"):
                doc.add_paragraph(line)
            doc.save(out)
        except ImportError:
            out = out.replace(".docx", ".txt")
            with open(out, "w", encoding="utf-8") as f:
                f.write(text)
    else:
        with open(out, "w", encoding="utf-8") as f:
            f.write(text)
    return out


_EXTRACTORS = {
    ".txt": _extract_text_from_txt,
    ".pdf": _extract_text_from_pdf,
    ".docx": _extract_text_from_docx,
}


def process_file(file_path: str, mode: str = "replace") -> dict:
    ext = Path(file_path).suffix.lower()

    if ext in (".csv", ".xlsx", ".xls"):
        out_path, entities = process_tabular(file_path, mode)
        by_type = {}
        for e in entities:
            by_type[e["label"]] = by_type.get(e["label"], 0) + 1
        return {
            "filename": os.path.basename(out_path),
            "entities": entities,
            "stats": {"total_entities": len(entities), "by_type": by_type},
            "download_url": f"/api/download/{os.path.basename(out_path)}",
        }

    if ext in (".png", ".jpg", ".jpeg", ".bmp", ".tiff"):
        try:
            import easyocr

            reader = easyocr.Reader(settings.ocr_languages, gpu=(settings.device != "cpu"))
            from app.services.image_pipeline import process_image

            out_path, entities = process_image(file_path, reader, [], mode)
        except ImportError:
            logger.warning("easyocr not installed, skipping OCR")
            entities = []

        by_type = {}
        for e in entities:
            by_type[e["label"]] = by_type.get(e["label"], 0) + 1
        return {
            "filename": os.path.basename(file_path).replace(
                Path(file_path).suffix, "_depersonalized.png"
            ),
            "entities": entities,
            "stats": {"total_entities": len(entities), "by_type": by_type},
            "download_url": "/api/download/"
            + os.path.basename(file_path).replace(Path(file_path).suffix, "_depersonalized.png"),
        }

    extractor = _EXTRACTORS.get(ext, _extract_text_from_txt)
    text = extractor(file_path)
    result = depersonalize_text(text, mode)

    out_path = _save_text_to_file(result["processed_text"], file_path)

    return {
        "filename": os.path.basename(out_path),
        "entities": result["entities"],
        "stats": result["stats"],
        "download_url": f"/api/download/{os.path.basename(out_path)}",
        "key": result.get("key"),
    }
